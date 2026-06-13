"""
Exportadores del reporte de auditoría de dependencias a TXT, Markdown, DOCX y PDF.

Todo es Python puro instalable con pip (python-docx, reportlab), sin depender de
Node, LibreOffice ni binarios del sistema — para que funcione tal cual en el
Windows del usuario. Cada formato refleja la misma estructura:

    Análisis General
    Análisis Técnico  (por aplicativo: versión + severidad, riesgos conocidos,
                       detalle expandido de CVE con enlaces, URL del desarrollador
                       y referencia de EOL, precisión técnica, mitigaciones)
    Conclusión General

DOCX y PDF llevan colores por nivel de riesgo, tablas con estilo y tipografías.
"""

import io
import re

# Paleta por nivel de riesgo (hex sin #)
_COLORS = {
    "crítico": "C0392B", "alto": "E67E22", "medio": "D4AC0D",
    "bajo": "7F8C8D", "ok": "27AE60", "desconocido": "8E44AD",
}
_ACCENT = "1F3A5F"      # azul corporativo para títulos
_HEADER_BG = "1F3A5F"
_LIGHT = "F2F4F7"


def _nivel_color(nivel):
    return _COLORS.get(nivel, "555555")


# ============================================================ TXT
def to_txt(rep):
    proj = rep["project"]; ov = rep["overall"]
    out = []
    line = "=" * 70
    out.append(line)
    out.append(f"  INFORME DE INVENTARIO DE SOFTWARE — {rep.get('generated_at','')}")
    out.append(f"  Proyecto: {proj.get('name') or 'n/d'}  v{proj.get('version') or 'n/d'}")
    out.append(f"  Generado por: {rep.get('author') or 'no indicado'}   |   Fecha y hora: {rep.get('generated_at','')}")
    out.append(f"  Recomendación sobre el inventario: {ov['verdict']}   |   Riesgo promedio: {ov['avg_score']}")
    out.append(f"  Insumo para la evaluación final de QA · {'con IA (Claude)' if rep.get('ai_used') else 'modelo de riesgo local'}")
    out.append(line)
    out.append("")
    out.append("ANÁLISIS GENERAL")
    out.append("-" * 70)
    out.append(_wrap(rep["general"]))
    out.append("")
    out.append("ANÁLISIS TÉCNICO")
    out.append("-" * 70)
    for t in rep["tecnico"]:
        out.append("")
        out.append(f"• {t['package']} {t['installed'] or ''}  [{t['severidad']}]")
        out.append(f"    Sección: {t['section']} | Última npm: {t['latest'] or 'n/d'}"
                   f" ({t.get('latest_release') or 's/f'}) | Atraso: {t['gap'] or 'al día'}")
        if t.get("homepage") or t.get("repository") or t.get("npm_url"):
            out.append(f"    Desarrollador: {t.get('homepage') or t.get('repository') or t.get('npm_url')}")
            out.append(f"    npm: {t.get('npm_url')}")
        if t.get("eol_url"):
            out.append(f"    EOL (ciclo de vida): {t['eol_url']}")
        elif t.get("eol_info"):
            out.append(f"    EOL: {t['eol_info']}")
        out.append("    Riesgos conocidos:")
        for r in t["riesgos_conocidos"]:
            out.append(f"      - {r}")
        # Detalle expandido de CVE
        for d in t.get("cve_detalle", []):
            out.append(f"      · {d.get('id')} (sev. {d.get('severity') or 'n/d'})"
                       f"{' - ' + d['summary'] if d.get('summary') else ''}")
            for u in (d.get("references") or [])[:3]:
                out.append(f"          ref: {u}")
        out.append(f"    Precisión técnica: {t['precision_tecnica']}")
        out.append("    Medidas de mitigación recomendadas:")
        for m in t["mitigaciones"]:
            out.append(f"      - {m}")
    out.append("")
    leg = rep.get("cumplimiento_legal") or []
    if leg:
        out.append("")
        out.append("ANÁLISIS DE CUMPLIMIENTO LEGAL (Ley 21.459 / 19.628)")
        out.append("-" * 70)
        for h in leg:
            out.append(f"• {h['package']}  [{h['nivel'].upper()}] — {h['categoria']}")
            out.append("    " + _wrap(h["rationale"], 66).replace("\n", "\n    "))
            out.append("    Disposiciones legales potencialmente aplicables:")
            for a in h.get("articulos", []):
                out.append(f"      - {a['art']}: {a['texto']}")
            out.append("    Recomendación: " + _wrap(h.get("recomendacion") or "", 60).replace("\n", "\n    "))
            out.append("    " + (h.get("disclaimer") or ""))
            out.append("")
    out.append("CONCLUSIÓN GENERAL")
    out.append("-" * 70)
    out.append(_wrap(rep["conclusion"]))
    out.append("")
    return "\n".join(out).encode("utf-8")


def _wrap(text, width=70):
    import textwrap
    return "\n".join(textwrap.fill(p, width) for p in str(text).split("\n"))


# ============================================================ Markdown
def to_md(rep):
    proj = rep["project"]; ov = rep["overall"]
    L = []
    L.append(f"# Informe de inventario de software — {rep.get('generated_at','')}")
    L.append("")
    L.append(f"**Proyecto:** {proj.get('name') or 'n/d'}  ·  **Versión:** {proj.get('version') or 'n/d'}  ")
    L.append(f"**Generado por:** {rep.get('author') or 'no indicado'}  ·  **Fecha y hora:** {rep.get('generated_at','')}  ")
    L.append(f"**Recomendación sobre el inventario:** {ov['verdict']}  ·  **Riesgo promedio:** {ov['avg_score']}  ")
    L.append(f"_Insumo para la evaluación final de QA. Generado {'con asistencia de IA (Claude)' if rep.get('ai_used') else 'por el modelo de riesgo local'}._")
    L.append("")
    c = ov["counts"]
    L.append(f"| Crítico | Alto | Medio | Bajo | OK |")
    L.append(f"|--------:|-----:|------:|-----:|---:|")
    L.append(f"| {c['crítico']} | {c['alto']} | {c['medio']} | {c['bajo']} | {c['ok']} |")
    L.append("")
    L.append("## Análisis General")
    L.append("")
    L.append(str(rep["general"]))
    L.append("")
    L.append("## Análisis Técnico")
    L.append("")
    for t in rep["tecnico"]:
        L.append(f"### {t['package']} `{t['installed'] or '—'}` — {t['severidad']}")
        L.append("")
        L.append(f"- **Sección:** {t['section']} · **Última npm:** {t['latest'] or 'n/d'} "
                 f"({t.get('latest_release') or 's/f'}) · **Atraso:** {t['gap'] or 'al día'}")
        links = []
        if t.get("homepage"): links.append(f"[sitio oficial]({t['homepage']})")
        if t.get("repository"): links.append(f"[repositorio]({t['repository']})")
        if t.get("npm_url"): links.append(f"[npm]({t['npm_url']})")
        if links:
            L.append(f"- **Desarrollador:** {' · '.join(links)}")
        if t.get("eol_url"):
            L.append(f"- **EOL / ciclo de vida:** [{t['eol_url']}]({t['eol_url']})")
        elif t.get("eol_info"):
            L.append(f"- **EOL:** {t['eol_info']}")
        L.append(f"- **Riesgos conocidos:**")
        for r in t["riesgos_conocidos"]:
            L.append(f"    - {r}")
        if t.get("cve_detalle"):
            L.append(f"- **Detalle de CVE:**")
            for d in t["cve_detalle"]:
                L.append(f"    - **{d.get('id')}** (severidad {d.get('severity') or 'n/d'})"
                         f"{' — ' + d['summary'] if d.get('summary') else ''}")
                for u in (d.get("references") or [])[:3]:
                    L.append(f"        - {u}")
        L.append(f"- **Precisión técnica:** {t['precision_tecnica']}")
        L.append(f"- **Medidas de mitigación recomendadas:**")
        for m in t["mitigaciones"]:
            L.append(f"    - {m}")
        L.append("")
    leg = rep.get("cumplimiento_legal") or []
    if leg:
        L.append("## Análisis de Cumplimiento Legal")
        L.append("")
        L.append("> Evaluación frente a la **Ley 21.459** (delitos informáticos) y la "
                 "**Ley 19.628** (protección de datos personales). Apoyo automatizado a la "
                 "revisión; no constituye asesoría legal.")
        L.append("")
        for h in leg:
            L.append(f"### ⚖ {h['package']} — nivel {h['nivel'].upper()} ({h['categoria']})")
            L.append("")
            L.append(h["rationale"])
            L.append("")
            L.append("**Disposiciones legales potencialmente aplicables:**")
            for a in h.get("articulos", []):
                L.append(f"- **{a['art']}** — {a['texto']}")
            L.append("")
            L.append(f"**Recomendación:** {h.get('recomendacion') or ''}")
            L.append("")
            L.append(f"_{h.get('disclaimer') or ''}_")
            L.append("")
    L.append("## Conclusión General")
    L.append("")
    L.append(str(rep["conclusion"]))
    L.append("")
    return "\n".join(L).encode("utf-8")


# ============================================================ DOCX
def to_docx(rep):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    proj = rep["project"]; ov = rep["overall"]
    doc = Document()
    # Estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)

    def set_cell(cell, text, bold=False, color=None, white=False, size=9.5):
        cell.text = ""
        p = cell.paragraphs[0]; run = p.add_run(str(text))
        run.bold = bold; run.font.size = Pt(size)
        if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif color: run.font.color.rgb = RGBColor.from_string(color)

    # ---- Portada / encabezado ----
    title = doc.add_paragraph()
    r = title.add_run(f"Informe de inventario de software")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string(_ACCENT)
    dt = title.add_run(f"   {rep.get('generated_at','')}")
    dt.font.size = Pt(12); dt.font.color.rgb = RGBColor.from_string("888888")
    sub = doc.add_paragraph()
    sr = sub.add_run(f"{proj.get('name') or 'proyecto'}  ·  versión {proj.get('version') or 'n/d'}")
    sr.font.size = Pt(11); sr.font.color.rgb = RGBColor.from_string("555555")
    au = doc.add_paragraph()
    ar = au.add_run(f"Generado por: {rep.get('author') or 'no indicado'}   ·   Fecha y hora: {rep.get('generated_at','')}")
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string("555555")

    vp = doc.add_paragraph()
    vr = vp.add_run(f"Recomendación sobre el inventario: {ov['verdict']}")
    vr.bold = True; vr.font.size = Pt(12)
    vr.font.color.rgb = RGBColor.from_string(_nivel_color(ov["verdict_level"]))
    doc.add_paragraph(f"Riesgo promedio: {ov['avg_score']}   ·   Insumo para la evaluación "
                      f"final de QA   ·   {'con asistencia de IA (Claude)' if rep.get('ai_used') else 'modelo de riesgo local'}")

    # tabla resumen de conteos
    c = ov["counts"]
    tb = doc.add_table(rows=2, cols=5); tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    heads = ["Crítico", "Alto", "Medio", "Bajo", "OK"]
    keys = ["crítico", "alto", "medio", "bajo", "ok"]
    for i, h in enumerate(heads):
        set_cell(tb.rows[0].cells[i], h, bold=True, white=True); shade(tb.rows[0].cells[i], _nivel_color(keys[i]))
        set_cell(tb.rows[1].cells[i], c.get(keys[i], 0), bold=True)
    tb.style = "Table Grid"

    # ---- Análisis General ----
    _h(doc, "Análisis General")
    doc.add_paragraph(str(rep["general"]))

    # ---- Análisis Técnico ----
    _h(doc, "Análisis Técnico")
    for t in rep["tecnico"]:
        hp = doc.add_paragraph()
        hr = hp.add_run(f"{t['package']}  {t['installed'] or '—'}")
        hr.bold = True; hr.font.size = Pt(12)
        sev = hp.add_run(f"   [{t['severidad']}]")
        sev.bold = True; sev.font.color.rgb = RGBColor.from_string(_nivel_color(t["nivel"]))

        meta = doc.add_paragraph()
        meta.add_run(f"Sección: {t['section']}  ·  Última npm: {t['latest'] or 'n/d'} "
                     f"({t.get('latest_release') or 's/f'})  ·  Atraso: {t['gap'] or 'al día'}").font.size = Pt(9)

        # Enlaces desarrollador + EOL
        link_p = doc.add_paragraph()
        link_p.add_run("Desarrollador: ").bold = True
        if t.get("homepage"): _hyperlink(link_p, "sitio oficial", t["homepage"]); link_p.add_run("  ")
        if t.get("repository"): _hyperlink(link_p, "repositorio", t["repository"]); link_p.add_run("  ")
        if t.get("npm_url"): _hyperlink(link_p, "npm", t["npm_url"])
        eol_p = doc.add_paragraph()
        eol_p.add_run("EOL / ciclo de vida: ").bold = True
        if t.get("eol_url"): _hyperlink(eol_p, t["eol_url"], t["eol_url"])
        else: eol_p.add_run(t.get("eol_info") or "n/d").font.size = Pt(9)

        _label(doc, "Riesgos conocidos")
        for rr in t["riesgos_conocidos"]:
            doc.add_paragraph(str(rr), style="List Bullet")
        if t.get("cve_detalle"):
            _label(doc, "Detalle de CVE")
            for d in t["cve_detalle"]:
                p = doc.add_paragraph(style="List Bullet")
                idr = p.add_run(f"{d.get('id')} "); idr.bold = True
                p.add_run(f"(severidad {d.get('severity') or 'n/d'})"
                          f"{' — ' + d['summary'] if d.get('summary') else ''}").font.size = Pt(9.5)
                for u in (d.get("references") or [])[:3]:
                    rp = doc.add_paragraph(style="List Bullet 2")
                    _hyperlink(rp, u, u)
        _label(doc, "Precisión técnica")
        doc.add_paragraph(str(t["precision_tecnica"]))
        _label(doc, "Medidas de mitigación recomendadas")
        for m in t["mitigaciones"]:
            doc.add_paragraph(str(m), style="List Bullet")
        doc.add_paragraph()

    # ---- Análisis de Cumplimiento Legal ----
    leg = rep.get("cumplimiento_legal") or []
    if leg:
        _h(doc, "Análisis de Cumplimiento Legal")
        intro = doc.add_paragraph()
        intro.add_run("Evaluación frente a la Ley 21.459 (delitos informáticos) y la "
                      "Ley 19.628 (protección de datos personales).").font.size = Pt(9)
        for h in leg:
            hp = doc.add_paragraph()
            hr = hp.add_run(f"⚖ {h['package']}  [{h['nivel'].upper()}]")
            hr.bold = True; hr.font.size = Pt(12)
            hr.font.color.rgb = RGBColor.from_string(_nivel_color(h["nivel"]))
            doc.add_paragraph(str(h["rationale"]))
            _label(doc, "Disposiciones legales potencialmente aplicables")
            for a in h.get("articulos", []):
                p = doc.add_paragraph(style="List Bullet")
                ar = p.add_run(f"{a['art']}: "); ar.bold = True
                p.add_run(str(a["texto"])).font.size = Pt(9.5)
            _label(doc, "Recomendación")
            doc.add_paragraph(str(h.get("recomendacion") or ""))
            dp = doc.add_paragraph()
            dr = dp.add_run(str(h.get("disclaimer") or ""))
            dr.italic = True; dr.font.size = Pt(8); dr.font.color.rgb = RGBColor.from_string("888888")
            doc.add_paragraph()

    # ---- Conclusión General ----
    _h(doc, "Conclusión General")
    doc.add_paragraph(str(rep["conclusion"]))

    buf = io.BytesIO(); doc.save(buf)
    return _fix_docx_zoom(buf.getvalue())


def _fix_docx_zoom(blob):
    """python-docx emits <w:zoom/> sin el atributo w:percent requerido por el
    esquema estricto. Word lo abre igual, pero lo añadimos para que el archivo
    sea estrictamente válido."""
    import zipfile
    try:
        src = io.BytesIO(blob)
        out = io.BytesIO()
        with zipfile.ZipFile(src) as zin:
            names = zin.namelist()
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
                for n in names:
                    data = zin.read(n)
                    if n == "word/settings.xml":
                        txt = data.decode("utf-8", "replace")
                        if "<w:zoom" in txt and "w:percent" not in txt.split("<w:zoom", 1)[1][:60]:
                            txt = re.sub(r"<w:zoom\s*/>", '<w:zoom w:percent="100"/>', txt)
                            txt = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*)>',
                                         r'<w:zoom w:percent="100"\1>', txt)
                        data = txt.encode("utf-8")
                    zout.writestr(n, data)
        return out.getvalue()
    except Exception:
        return blob


def _h(doc, text):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph(); p.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(_ACCENT)
    # borde inferior
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", _ACCENT)):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)


def _label(doc, text):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph(); r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string("777777")


def _hyperlink(paragraph, text, url):
    """Insert a real clickable hyperlink run into a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1155CC"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)


def _table_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl; tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        for k, v in (("w:val", "single"), ("w:sz", "4"), ("w:space", "0"), ("w:color", "CCCCCC")):
            el.set(qn(k), v)
        borders.append(el)
    tblPr.append(borders)


# ============================================================ PDF
def to_pdf(rep):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, ListFlowable, ListItem, HRFlowable)
    from reportlab.lib.enums import TA_LEFT

    proj = rep["project"]; ov = rep["overall"]
    buf = io.BytesIO()
    docp = SimpleDocTemplate(buf, pagesize=letter, topMargin=18*mm, bottomMargin=18*mm,
                             leftMargin=18*mm, rightMargin=18*mm,
                             title=f"Reporte {proj.get('name') or ''}")
    ss = getSampleStyleSheet()
    accent = colors.HexColor("#" + _ACCENT)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], textColor=accent, fontSize=18, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], textColor=accent, fontSize=13,
                        spaceBefore=12, spaceAfter=6, borderWidth=0)
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=14)
    small = ParagraphStyle("small", parent=body, fontSize=8, textColor=colors.HexColor("#666666"))
    pkgst = ParagraphStyle("pkg", parent=body, fontSize=11.5, spaceBefore=8)
    labelst = ParagraphStyle("lbl", parent=body, fontSize=7.5, textColor=colors.HexColor("#777777"))

    def esc(s):
        return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = []
    story.append(Paragraph(f"Informe de inventario de software &nbsp;<font size='12' color='#888888'>{esc(rep.get('generated_at',''))}</font>", h1))
    story.append(Paragraph(f"{esc(proj.get('name') or 'proyecto')} &nbsp;·&nbsp; versión {esc(proj.get('version') or 'n/d')}", small))
    story.append(Paragraph(f"Generado por: {esc(rep.get('author') or 'no indicado')} &nbsp;·&nbsp; Fecha y hora: {esc(rep.get('generated_at',''))}", small))
    story.append(Spacer(1, 4))
    vcol = colors.HexColor("#" + _nivel_color(ov["verdict_level"]))
    story.append(Paragraph(f'<b><font color="#{_nivel_color(ov["verdict_level"])}">Recomendación sobre el inventario: {esc(ov["verdict"])}</font></b> '
                           f'&nbsp; Riesgo promedio: {ov["avg_score"]}', body))
    story.append(Paragraph("Insumo para la evaluación final de QA · " + ("con asistencia de IA (Claude)" if rep.get("ai_used")
                           else "modelo de riesgo local"), small))
    story.append(Spacer(1, 6))

    # tabla resumen
    c = ov["counts"]
    data = [["Crítico", "Alto", "Medio", "Bajo", "OK"],
            [c["crítico"], c["alto"], c["medio"], c["bajo"], c["ok"]]]
    tb = Table(data, colWidths=[34*mm]*5)
    tstyle = [("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
              ("ALIGN", (0, 0), (-1, -1), "CENTER"),
              ("FONTSIZE", (0, 0), (-1, -1), 9),
              ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
              ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")]
    for i, k in enumerate(["crítico", "alto", "medio", "bajo", "ok"]):
        tstyle.append(("BACKGROUND", (i, 0), (i, 0), colors.HexColor("#" + _nivel_color(k))))
    tb.setStyle(TableStyle(tstyle))
    story.append(tb)

    story.append(Paragraph("Análisis General", h2))
    story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))
    for para in str(rep["general"]).split("\n"):
        if para.strip():
            story.append(Paragraph(esc(para), body))

    story.append(Paragraph("Análisis Técnico", h2))
    story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))
    for t in rep["tecnico"]:
        lvlcol = _nivel_color(t["nivel"])
        story.append(Paragraph(
            f'<b>{esc(t["package"])} {esc(t["installed"] or "—")}</b> '
            f'&nbsp;<font color="#{lvlcol}"><b>[{esc(t["severidad"])}]</b></font>', pkgst))
        story.append(Paragraph(
            f'Sección: {esc(t["section"])} · Última npm: {esc(t["latest"] or "n/d")} '
            f'({esc(t.get("latest_release") or "s/f")}) · Atraso: {esc(t["gap"] or "al día")}', small))
        # enlaces
        links = []
        if t.get("homepage"): links.append(f'<a href="{esc(t["homepage"])}" color="#1155CC">sitio oficial</a>')
        if t.get("repository"): links.append(f'<a href="{esc(t["repository"])}" color="#1155CC">repositorio</a>')
        if t.get("npm_url"): links.append(f'<a href="{esc(t["npm_url"])}" color="#1155CC">npm</a>')
        if links:
            story.append(Paragraph("Desarrollador: " + " · ".join(links), small))
        if t.get("eol_url"):
            story.append(Paragraph(f'EOL / ciclo de vida: <a href="{esc(t["eol_url"])}" color="#1155CC">{esc(t["eol_url"])}</a>', small))
        elif t.get("eol_info"):
            story.append(Paragraph("EOL: " + esc(t["eol_info"]), small))

        story.append(Paragraph("RIESGOS CONOCIDOS", labelst))
        story.append(ListFlowable([ListItem(Paragraph(esc(r), body), leftIndent=10)
                                   for r in t["riesgos_conocidos"]], bulletType="bullet", start="•"))
        if t.get("cve_detalle"):
            story.append(Paragraph("DETALLE DE CVE", labelst))
            items = []
            for d in t["cve_detalle"]:
                txt = f'<b>{esc(d.get("id"))}</b> (severidad {esc(d.get("severity") or "n/d")})'
                if d.get("summary"): txt += " — " + esc(d["summary"])
                refs = (d.get("references") or [])[:3]
                if refs:
                    txt += "<br/>" + "<br/>".join(f'<a href="{esc(u)}" color="#1155CC">{esc(u)}</a>' for u in refs)
                items.append(ListItem(Paragraph(txt, body), leftIndent=10))
            story.append(ListFlowable(items, bulletType="bullet", start="•"))
        story.append(Paragraph("PRECISIÓN TÉCNICA", labelst))
        story.append(Paragraph(esc(t["precision_tecnica"]), body))
        story.append(Paragraph("MEDIDAS DE MITIGACIÓN RECOMENDADAS", labelst))
        story.append(ListFlowable([ListItem(Paragraph(esc(m), body), leftIndent=10)
                                   for m in t["mitigaciones"]], bulletType="bullet", start="•"))
        story.append(Spacer(1, 4))

    leg = rep.get("cumplimiento_legal") or []
    if leg:
        story.append(Paragraph("Análisis de Cumplimiento Legal", h2))
        story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))
        story.append(Paragraph("Evaluación frente a la Ley 21.459 (delitos informáticos) y la "
                               "Ley 19.628 (protección de datos personales). Apoyo a la revisión; "
                               "no constituye asesoría legal.", small))
        for h in leg:
            lc = _nivel_color(h["nivel"])
            story.append(Paragraph(
                f'<b><font color="#{lc}">⚖ {esc(h["package"])} [{esc(h["nivel"].upper())}]</font></b> '
                f'<font size="8">— {esc(h["categoria"])}</font>', pkgst))
            story.append(Paragraph(esc(h["rationale"]), body))
            story.append(Paragraph("DISPOSICIONES LEGALES POTENCIALMENTE APLICABLES", labelst))
            story.append(ListFlowable(
                [ListItem(Paragraph(f"<b>{esc(a['art'])}:</b> {esc(a['texto'])}", body), leftIndent=10)
                 for a in h.get("articulos", [])], bulletType="bullet", start="•"))
            story.append(Paragraph("RECOMENDACIÓN", labelst))
            story.append(Paragraph(esc(h.get("recomendacion") or ""), body))
            story.append(Paragraph(f'<i><font color="#888888" size="7">{esc(h.get("disclaimer") or "")}</font></i>', small))
            story.append(Spacer(1, 4))

    story.append(Paragraph("Conclusión General", h2))
    story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))
    for para in str(rep["conclusion"]).split("\n"):
        if para.strip():
            story.append(Paragraph(esc(para), body))

    docp.build(story)
    return buf.getvalue()


# ============================================================ CSV
def to_csv(rep):
    """CSV plano por dependencia, para abrir en Excel (con BOM UTF-8)."""
    import csv as _csv
    sio = io.StringIO()
    w = _csv.writer(sio)
    proj = rep.get("project") or {}
    ov = rep.get("overall") or {}
    w.writerow([f"# Informe de inventario de software - {rep.get('generated_at','')}"])
    w.writerow([f"# Proyecto: {proj.get('name') or 'n/d'}  version: {proj.get('version') or 'n/d'}"])
    w.writerow([f"# Generado por: {rep.get('author') or 'no indicado'}  Fecha y hora: {rep.get('generated_at','')}"])
    w.writerow([f"# Recomendacion sobre el inventario: {ov.get('verdict')}  riesgo_promedio: {ov.get('avg_score')}  (insumo para evaluacion final de QA)"])
    w.writerow([])
    w.writerow(["paquete", "seccion", "instalada", "ultima_npm", "release_ultima",
                "atraso", "nivel_riesgo", "severidad", "cve_count", "cve_ids",
                "riesgos_conocidos", "precision_tecnica", "mitigaciones",
                "sitio_oficial", "repositorio", "npm_url", "eol_url",
                "cumplimiento_nivel", "cumplimiento_categoria", "cumplimiento_articulos"])
    for t in rep.get("tecnico", []):
        comp = t.get("cumplimiento") or {}
        arts = "; ".join(f"{a['art']}: {a['texto']}" for a in comp.get("articulos", []))
        w.writerow([
            t.get("package"), t.get("section"), t.get("installed") or "",
            t.get("latest") or "", t.get("latest_release") or "",
            t.get("gap") or "al día", t.get("nivel"), t.get("severidad"),
            t.get("vuln_count") or 0, " ".join(t.get("vuln_ids") or []),
            " | ".join(t.get("riesgos_conocidos") or []),
            t.get("precision_tecnica") or "",
            " | ".join(t.get("mitigaciones") or []),
            t.get("homepage") or "", t.get("repository") or "",
            t.get("npm_url") or "", t.get("eol_url") or "",
            comp.get("nivel") or "", comp.get("categoria") or "", arts,
        ])
    return ("\ufeff" + sio.getvalue()).encode("utf-8")


# ============================================================ dispatch
EXPORTERS = {"txt": (to_txt, "text/plain; charset=utf-8"),
             "md":  (to_md,  "text/markdown; charset=utf-8"),
             "csv": (to_csv, "text/csv; charset=utf-8"),
             "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
             "pdf": (to_pdf, "application/pdf")}


def export(rep, fmt):
    fmt = (fmt or "md").lower()
    if fmt not in EXPORTERS:
        raise ValueError(f"formato no soportado: {fmt}")
    fn, mime = EXPORTERS[fmt]
    return fn(rep), mime


# ============================================================================
# Informe de cumplimiento (análisis estático de código fuente)
# Formato: carta, Arial 11, justificado, interlineado 1.15, alto contraste.
# ============================================================================
_SEV_COLOR = {"crítico": "B00020", "alto": "C05600", "medio": "8A6D00",
              "bajo": "1E7A34", "info": "3A4A5A"}  # tonos oscuros: contraste sobre blanco
_SEV_ORDER = {"crítico": 4, "alto": 3, "medio": 2, "bajo": 1, "info": 0}
# Colores por NIVEL DE RIESGO contextual (impacto × probabilidad).
_NR_COLOR = {"Crítico": "B00020", "Alto": "C05600", "Medio": "8A6D00",
             "Bajo": "1E7A34", "Informativo": "3A4A5A"}
_MARCO = ("Guía Técnica «Lineamientos para el Desarrollo de Software», División de "
          "Gobierno Digital (SEGPRES) v2.0 (2021); Ley 21.180 de Transformación Digital "
          "del Estado y sus normas técnicas (Decretos 7, 9, 10 y 11 de 2023, SEGPRES); "
          "DS 83/2004 (seguridad del documento electrónico); Leyes 19.628, 19.799, "
          "21.096 y 21.459. Referencias técnicas: OWASP MASVS/ASVS y CWE.")

# Etiquetas legibles y colores por estado de control de cumplimiento.
_ESTADO_LABEL = {"cumple": "Cumple", "no_cumple": "No cumple",
                 "observado": "Observado", "no_evaluable": "No evaluable"}
_ESTADO_COLOR = {"cumple": "1E7A34", "no_cumple": "B00020",
                 "observado": "8A6D00", "no_evaluable": "3A4A5A"}


def _cs_sorted(findings):
    # Si vienen con 'score' (razonamiento), ordenar por score; si no, por severidad.
    if findings and "score" in findings[0]:
        return sorted(findings, key=lambda f: -f.get("score", 0))
    return sorted(findings, key=lambda f: (-_SEV_ORDER.get(f["severidad"], 0),
                                           f["categoria"], f["archivo"]))


def _cs_dedup_analisis(f):
    """Devuelve el texto de 'análisis' evitando que repita literalmente la guía
    y el riesgo de omisión (que ya se listan en campos propios)."""
    an = (f.get("analisis") or "").strip()
    guia = (f.get("guia") or "").strip()
    rom = (f.get("riesgo_omision") or "").strip()
    # Si el análisis es solo la concatenación de título+guía+riesgo (fallback sin
    # IA), lo dejamos vacío para no duplicar; el lector verá guía y riesgo aparte.
    if an and (guia and guia in an) and (rom and rom in an):
        return ""
    return an


def codescan_to_md(report, author="no indicado", generated_at=""):
    R = report
    st = R.get("stats", {}); rs = R.get("resumen", {}); an = R.get("analisis", {})
    arch = R.get("archive", {}); inv = R.get("inventario", {})
    cmp = R.get("cumplimiento") or {}
    L = []
    L.append("# Informe de auditoría de cumplimiento y aseguramiento de calidad")
    L.append("## Análisis estático de seguridad de código (SAST) y conformidad normativa")
    L.append("")
    L.append(f"**Artefacto analizado:** {arch.get('nombre','(s/d)')}  ·  "
             f"{arch.get('files','?')} archivos  ·  {round(arch.get('bytes',0)/1048576,1)} MB")
    L.append(f"**Generado por:** {author}  ·  **Fecha y hora:** {generated_at}")
    L.append(f"**Clasificación del informe:** Uso interno — insumo para evaluación de QA y auditoría")
    L.append("")
    L.append(f"**Marco de referencia:** {_MARCO}")
    L.append("")
    L.append("> Análisis estático automatizado con apoyo de razonamiento (aprendizaje "
             "automático local y, cuando está disponible, modelo de lenguaje). Constituye un "
             "insumo para la evaluación de aseguramiento de calidad y la auditoría de "
             "cumplimiento; no es un dictamen jurídico ni una certificación de conformidad. "
             "La evidencia de secretos se presenta enmascarada por el principio de mínima exposición.")
    L.append("")

    # 1) RESUMEN EJECUTIVO
    L.append("## 1. Resumen ejecutivo")
    L.append("")
    if an.get("resumen_ejecutivo"):
        L.append(an["resumen_ejecutivo"]); L.append("")
    if cmp:
        L.append(f"**Veredicto de cumplimiento normativo:** {cmp.get('veredicto','—')} "
                 f"({cmp.get('pct_cumplimiento',0)}% de controles evaluables conformes). "
                 f"{cmp.get('veredicto_detalle','')}")
        L.append("")

    # 2) VEREDICTO Y MATRIZ DE CUMPLIMIENTO
    if cmp:
        co = cmp.get("conteos", {})
        L.append("## 2. Conformidad con el marco normativo")
        L.append("")
        L.append(f"Se evaluaron **{cmp.get('total_controles',0)} controles** auditables "
                 f"derivados de la normativa aplicable: "
                 f"**{co.get('cumple',0)} conformes**, **{co.get('no_cumple',0)} no conformes**, "
                 f"**{co.get('observado',0)} observados** y **{co.get('no_evaluable',0)} no "
                 f"evaluables por análisis estático** (controles organizacionales que requieren "
                 f"verificación documental).")
        L.append("")
        L.append("### Matriz de conformidad por instrumento")
        L.append("")
        L.append("| Instrumento | Cumple | No cumple | Observado | No evaluable |")
        L.append("|---|---:|---:|---:|---:|")
        for inst in cmp.get("instrumentos", []):
            iid = inst["id"]
            d = cmp.get("por_instrumento", {}).get(iid, {})
            L.append(f"| {iid} — {inst['titulo']} | {d.get('cumple',0)} | "
                     f"{d.get('no_cumple',0)} | {d.get('observado',0)} | {d.get('no_evaluable',0)} |")
        L.append("")
        # Funciones del Decreto 7 (marco tipo CSF)
        pf = cmp.get("por_funcion", {})
        if any(sum(v.values()) for v in pf.values()):
            L.append("### Cobertura por función de ciberseguridad (Decreto 7/2023)")
            L.append("")
            L.append("| Función | Cumple | No cumple | Observado | No evaluable |")
            L.append("|---|---:|---:|---:|---:|")
            for fn, d in pf.items():
                if sum(d.values()):
                    L.append(f"| {fn} | {d.get('cumple',0)} | {d.get('no_cumple',0)} | "
                             f"{d.get('observado',0)} | {d.get('no_evaluable',0)} |")
            L.append("")

        # 3) CONTROLES NO CONFORMES (detalle)
        no_conf = [c for c in cmp.get("controles", []) if c["estado"] == "no_cumple"]
        if no_conf:
            L.append("### Controles no conformes (requieren remediación)")
            L.append("")
            for c in no_conf:
                L.append(f"- **{c['instrumento']} {c['referencia']} — {c['titulo']}** "
                         f"_(severidad {c['severidad']})_")
                L.append(f"  - Exigencia: {c['exigencia']}")
                if c.get("hallazgos"):
                    refs = "; ".join(
                        f"`{h['rule_id']}` en {str(h.get('archivo','')).split('/')[-1]}"
                        + (f":{h['linea']}" if h.get("linea") else "")
                        for h in c["hallazgos"][:5])
                    L.append(f"  - Evidencia ({c['n_hallazgos']}): {refs}")
            L.append("")

        # Controles no evaluables (transparencia de alcance)
        no_eval = [c for c in cmp.get("controles", []) if c["estado"] == "no_evaluable"]
        if no_eval:
            L.append("### Controles no evaluables por análisis estático")
            L.append("")
            L.append("Los siguientes controles son de naturaleza organizacional o procedimental "
                     "y deben acreditarse mediante evidencia documental independiente:")
            L.append("")
            for c in no_eval:
                doc = f" Evidencia esperada: {c['evidencia_documental']}" if c.get("evidencia_documental") else ""
                L.append(f"- **{c['instrumento']} {c['referencia']} — {c['titulo']}.**{doc}")
            L.append("")

    # 4) PANORAMA TÉCNICO
    L.append("## 3. Panorama técnico del análisis estático")
    L.append("")
    L.append(f"- Total de observaciones: **{rs.get('total',0)}**  ·  nivel de riesgo agregado: "
             f"**{an.get('nivel_global','—')}**")
    sev = rs.get("por_severidad", {})
    L.append("- Por severidad: " + ", ".join(f"{k}: {v}" for k, v in
             sorted(sev.items(), key=lambda x: -_SEV_ORDER.get(x[0], 0))))
    L.append("- Por ámbito: " + ", ".join(f"{k}: {v}" for k, v in rs.get("por_categoria", {}).items()))
    L.append("- Lenguajes: " + (", ".join(f"{k} ({v})" for k, v in (st.get('lenguajes') or {}).items()) or "—"))
    if inv.get("agregados"):
        L.append(f"- Componentes integrados al inventario de software: **{inv['agregados']}**")
    if st.get("permisos_android"):
        L.append(f"- Permisos Android declarados ({len(st['permisos_android'])}): "
                 + ", ".join(p.split('.')[-1] for p in st['permisos_android']))
    if an.get("hotspots"):
        L.append("- Puntos calientes (archivos con múltiples hallazgos): "
                 + "; ".join(f"{h['archivo'].split('/')[-1]} ({len(h['reglas'])})" for h in an['hotspots'][:6]))
    # Motores de análisis (defensa en profundidad).
    mot = (R.get("motores") or {}).get("por_motor") or rs.get("por_motor") or {}
    if mot:
        L.append("- Motores de análisis (hallazgos por motor): "
                 + ", ".join(f"{k}: {v}" for k, v in mot.items()))
    corro = [f for f in R.get("findings", []) if f.get("corroborado_por")]
    if corro:
        L.append(f"- Hallazgos corroborados por más de un motor: **{len(corro)}** "
                 "(mayor confianza)")
    # Análisis de flujo de datos (taint).
    taint = R.get("_taint") or []
    if taint:
        L.append(f"- Rutas de flujo de datos no confiable detectadas (fuente → sumidero): "
                 f"**{len(taint)}**")
        for tp in taint[:6]:
            L.append(f"  - {tp['archivo'].split('/')[-1]}: {tp['fuente']} (L{tp['fuente_linea']}) "
                     f"→ {tp['sumidero']} (L{tp['sumidero_linea']}) · {tp.get('cwe','')}")
    L.append("")

    # 5) ANÁLISIS TÉCNICO DETALLADO (separado del cumplimiento normativo/legal)
    L.append("## 4. Análisis técnico detallado")
    L.append("")
    L.append("Esta sección desarrolla cada observación en el contexto real de uso del código "
             "(inteligencia local de la herramienta). El nivel de riesgo se calcula como "
             "**impacto × probabilidad**; ambos factores se justifican. El mapeo a la normativa "
             "y la ley se trata por separado en la sección 2.")
    L.append("")
    for i, f in enumerate(_cs_sorted(R.get("findings", [])), 1):
        ctx = f.get("contexto") or {}
        nr = f.get("nivel_riesgo", "—")
        L.append(f"### 4.{i} [{nr}] {f['titulo']}  (`{f['rule_id']}`"
                 + (f" · {f['cwe']}" if f.get("cwe") else "") + ")")
        L.append("")
        if f.get("descripcion_contextual"):
            L.append(f"**Descripción (contexto detectado).** {f['descripcion_contextual']}")
            L.append("")
        if f.get("explicacion_no_tecnica"):
            L.append(f"**Qué significa, en simple.** {f['explicacion_no_tecnica']}")
            L.append("")
        loc = f"{f['archivo']}" + (f":{f['linea']}" if f.get("linea") else "")
        L.append(f"**Ubicación.** `{loc}` · rol del archivo: {ctx.get('rol_archivo','—')}"
                 + (f" · {ctx.get('constructo')}" if ctx.get("constructo") else "")
                 + f" · exposición: {f.get('exposicion','—')}")
        L.append("")
        # Procedencia (motor) y corroboración entre motores.
        motor = f.get("motor", "interno")
        proc = f"motor: {motor}"
        if f.get("corroborado_por"):
            proc += f"; corroborado por: {', '.join(f['corroborado_por'])}"
        if f.get("secreto_prob") is not None:
            proc += (f"; verosimilitud de secreto (ML): {f['secreto_prob']:.0%}"
                     + (" — posible falso positivo" if f.get("posible_falso_positivo") else ""))
        L.append(f"**Procedencia y verificación.** {proc}.")
        L.append("")
        # Explotabilidad por flujo de datos (taint).
        if f.get("alcanzable_taint") and f.get("ruta_flujo"):
            rf = f["ruta_flujo"]
            L.append(f"**Explotabilidad (flujo de datos).** Ruta detectada: {rf['fuente']} "
                     f"(línea {rf['fuente_linea']}) → {rf['sumidero']} (línea {rf['sumidero_linea']}), "
                     f"a través de la variable «{rf['var']}». El dato no confiable alcanza la "
                     "operación sensible, por lo que la observación se considera potencialmente "
                     "explotable.")
            L.append("")
        frag = ctx.get("fragmento") or []
        if frag:
            L.append("**Evidencia en contexto.**")
            L.append("")
            L.append("```")
            for (ln, txt, es) in frag:
                L.append(f"{'>' if es else ' '} {ln:>4} | {txt}")
            L.append("```")
            L.append("")
        else:
            L.append(f"**Evidencia (extracto).** `{f['evidencia']}`")
            L.append("")
        # Cálculo del nivel de riesgo (impacto × probabilidad), con justificación.
        L.append(f"**Cálculo del nivel de riesgo.** Severidad intrínseca de la regla: "
                 f"{f.get('severidad','—')}.")
        L.append(f"- Impacto: **{f.get('impacto_nivel','—')}** ({f.get('impacto','?')}/5). "
                 f"{f.get('impacto_just','')}")
        L.append(f"- Probabilidad: **{f.get('probabilidad_nivel','—')}** ({f.get('probabilidad','?')}/5). "
                 f"{f.get('probabilidad_just','')}")
        L.append(f"- Resultado: {f.get('impacto','?')} × {f.get('probabilidad','?')} = "
                 f"{f.get('riesgo_valor','?')} → nivel de riesgo **{nr}**.")
        L.append("")
        if f.get("impacto_contexto"):
            L.append(f"**Impacto en el contexto del código.** {f['impacto_contexto']}")
            L.append("")
        L.append(f"**Mitigación puntual.** {f.get('mitigacion_puntual', f.get('mitigacion',''))}")
        L.append("")
        if f.get("controles"):
            L.append(f"_Trazabilidad de cumplimiento (ver sección 2): {', '.join(f['controles'])}._")
            L.append("")

    # 6) MARCO NORMATIVO DE REFERENCIA
    if cmp.get("instrumentos"):
        L.append("## 5. Marco normativo de referencia")
        L.append("")
        for inst in cmp["instrumentos"]:
            L.append(f"- **{inst['id']}** — {inst['titulo']} ({inst['ley']}, {inst['fecha']}). "
                     f"{inst['url']}")
        L.append("")

    # 7) ANEXO METODOLÓGICO
    L.append("## 6. Anexo metodológico")
    L.append("")
    L.append("El análisis se realiza mediante inspección estática de patrones (SAST ligero) sobre "
             "el código fuente extraído en un entorno aislado, sin ejecutarlo. Cada regla mapea a "
             "una cláusula de la Guía SEGPRES, la norma legal aplicable, el control normativo "
             "auditable y, cuando corresponde, el identificador CWE. El motor de cumplimiento cruza "
             "los hallazgos con un catálogo de controles derivados artículo por artículo de la "
             "normativa, emitiendo por control uno de cuatro estados: cumple, no cumple, observado "
             "o no evaluable. Los controles organizacionales se reportan como no evaluables con "
             "indicación de la evidencia documental que los acreditaría.")
    L.append("")
    L.append("**Análisis avanzado (robustecimiento).** El alcance se refuerza con tres técnicas "
             "complementarias: (1) motores SAST de código abierto de la industria —Semgrep "
             "(multi-lenguaje, con reglas locales offline), detect-secrets y Bandit— cuyos "
             "hallazgos se normalizan e integran; cuando dos motores coinciden, el hallazgo se "
             "marca como corroborado (mayor confianza). (2) Análisis de flujo de datos (taint) "
             "intra-archivo, que rastrea entradas no confiables (fuentes) hasta operaciones "
             "sensibles (sumideros) para identificar rutas potencialmente explotables; esto "
             "aproxima el razonamiento de ejecución sin ejecutar el código. (3) Un clasificador de "
             "verosimilitud de secretos (regresión logística) que ajusta la probabilidad y "
             "distingue secretos reales de marcadores de posición.")
    L.append("")
    L.append("**Sobre las pruebas dinámicas (DAST).** Por seguridad, la herramienta NUNCA ejecuta "
             "el código auditado: ejecutar artefactos no confiables expondría al evaluador. Por "
             "ello no se realiza DAST en sentido estricto (ejecución del sistema en marcha); el "
             "componente \"dinámico\" se cubre mediante el análisis de flujo de datos y el "
             "razonamiento de rutas de ataque asistido por IA. Para pruebas dinámicas completas "
             "(con la aplicación efectivamente desplegada en un entorno controlado) se recomienda "
             "complementar con un proceso de DAST dedicado.")
    L.append("")
    if cmp.get("disclaimer"):
        L.append(f"> {cmp['disclaimer']}")
        L.append("")
    return "\n".join(L).encode("utf-8")


def _cs_set_doc_defaults(doc):
    """Arial 11, interlineado 1.0 (sencillo) SIN espacio entre párrafos,
    justificado, color de alto contraste."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string("1A1A1A")  # casi negro: alto contraste
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE      # interlineado 1.0
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(0)                             # sin espacio entre líneas/párrafos
    pf.space_before = Pt(0)
    # Encabezados: interlineado 1.0; mínima separación de sección por legibilidad.
    for nm in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            hpf = doc.styles[nm].paragraph_format
            hpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            hpf.line_spacing = 1.0
            hpf.space_after = Pt(0)
            hpf.space_before = Pt(6) if nm == "Heading 1" else Pt(4)
        except Exception:
            pass


def _cs_letter_page(doc):
    """Tamaño carta (Letter) con márgenes de 1 pulgada."""
    from docx.shared import Inches
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def _cs_shade_cell(cell, hex_color):
    """Sombrea una celda de tabla (fondo) en DOCX."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cs_topborder(paragraph, color="B0B7C3", size="6"):
    """Agrega un borde superior al párrafo: separa observaciones sin insertar
    líneas en blanco (respeta el interlineado 1.0 sin espacio entre párrafos)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    pbdr.append(top)
    pPr.append(pbdr)


def codescan_to_docx(report, author="no indicado", generated_at=""):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    R = report
    st = R.get("stats", {}); rs = R.get("resumen", {}); an = R.get("analisis", {})
    arch = R.get("archive", {}); inv = R.get("inventario", {})
    cmp = R.get("cumplimiento") or {}
    doc = Document()
    _cs_letter_page(doc)
    _cs_set_doc_defaults(doc)

    # ---------- PORTADA ----------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("Informe de auditoría de cumplimiento\ny aseguramiento de calidad")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string(_ACCENT)
    sub = doc.add_paragraph()
    sr = sub.add_run("Análisis estático de seguridad de código (SAST) y conformidad normativa")
    sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = RGBColor.from_string("333333")
    doc.add_paragraph("")
    for txt in (f"Artefacto analizado: {arch.get('nombre','(s/d)')}  ·  {arch.get('files','?')} archivos  ·  "
                f"{round(arch.get('bytes',0)/1048576,1)} MB",
                f"Generado por: {author}",
                f"Fecha y hora: {generated_at}",
                "Clasificación: Uso interno — insumo para evaluación de QA y auditoría"):
        p = doc.add_paragraph(); p.add_run(txt).font.size = Pt(11)
    mp = doc.add_paragraph(); mr = mp.add_run("Marco de referencia: " + _MARCO)
    mr.font.size = Pt(9); mr.font.color.rgb = RGBColor.from_string("333333")
    nota = doc.add_paragraph()
    nr = nota.add_run("Análisis estático automatizado con razonamiento (aprendizaje automático "
                      "local y modelo de lenguaje cuando está disponible): insumo para la "
                      "evaluación de aseguramiento de calidad y la auditoría de cumplimiento, no "
                      "un dictamen jurídico ni una certificación. La evidencia de secretos se "
                      "muestra enmascarada (mínima exposición).")
    nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = RGBColor.from_string("444444")

    # ---------- 1. RESUMEN EJECUTIVO ----------
    doc.add_heading("1. Resumen ejecutivo", level=1)
    if an.get("resumen_ejecutivo"):
        p = doc.add_paragraph(an["resumen_ejecutivo"]); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if cmp:
        p = doc.add_paragraph()
        lr = p.add_run("Veredicto de cumplimiento normativo: "); lr.bold = True
        vr = p.add_run(f"{cmp.get('veredicto','—')} ")
        vr.bold = True
        vc = "1E7A34" if cmp.get("veredicto") == "CONFORME" else (
            "B00020" if cmp.get("veredicto") == "NO CONFORME" else "8A6D00")
        vr.font.color.rgb = RGBColor.from_string(vc)
        p.add_run(f"({cmp.get('pct_cumplimiento',0)}% de controles evaluables conformes). "
                  f"{cmp.get('veredicto_detalle','')}")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---------- 2. CONFORMIDAD NORMATIVA ----------
    if cmp:
        co = cmp.get("conteos", {})
        doc.add_heading("2. Conformidad con el marco normativo", level=1)
        doc.add_paragraph(
            f"Se evaluaron {cmp.get('total_controles',0)} controles auditables derivados de la "
            f"normativa aplicable: {co.get('cumple',0)} conformes, {co.get('no_cumple',0)} no "
            f"conformes, {co.get('observado',0)} observados y {co.get('no_evaluable',0)} no "
            f"evaluables por análisis estático (controles organizacionales que requieren "
            f"verificación documental).")

        # Matriz por instrumento
        doc.add_heading("Matriz de conformidad por instrumento", level=2)
        instrumentos = cmp.get("instrumentos", [])
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for j, h in enumerate(("Instrumento", "Cumple", "No cumple", "Observado", "No eval.")):
            hdr[j].text = h
            for pp in hdr[j].paragraphs:
                for rr in pp.runs:
                    rr.bold = True; rr.font.size = Pt(10)
            _cs_shade_cell(hdr[j], _ACCENT)
            for pp in hdr[j].paragraphs:
                for rr in pp.runs:
                    rr.font.color.rgb = RGBColor.from_string("FFFFFF")
        for inst in instrumentos:
            d = cmp.get("por_instrumento", {}).get(inst["id"], {})
            row = tbl.add_row().cells
            row[0].text = f"{inst['id']} — {inst['titulo']}"
            row[1].text = str(d.get("cumple", 0)); row[2].text = str(d.get("no_cumple", 0))
            row[3].text = str(d.get("observado", 0)); row[4].text = str(d.get("no_evaluable", 0))
            for c in row:
                for pp in c.paragraphs:
                    for rr in pp.runs:
                        rr.font.size = Pt(9)
            if d.get("no_cumple", 0):
                _cs_shade_cell(row[2], "F4CCCC")
        doc.add_paragraph("")

        # Cobertura por función Decreto 7
        pf = cmp.get("por_funcion", {})
        if any(sum(v.values()) for v in pf.values()):
            doc.add_heading("Cobertura por función de ciberseguridad (Decreto 7/2023)", level=2)
            tf = doc.add_table(rows=1, cols=5); tf.style = "Light Grid Accent 1"
            h = tf.rows[0].cells
            for j, txt in enumerate(("Función", "Cumple", "No cumple", "Observado", "No eval.")):
                h[j].text = txt
                _cs_shade_cell(h[j], _ACCENT)
                for pp in h[j].paragraphs:
                    for rr in pp.runs:
                        rr.bold = True; rr.font.size = Pt(10)
                        rr.font.color.rgb = RGBColor.from_string("FFFFFF")
            for fn, d in pf.items():
                if sum(d.values()):
                    row = tf.add_row().cells
                    row[0].text = fn
                    row[1].text = str(d.get("cumple", 0)); row[2].text = str(d.get("no_cumple", 0))
                    row[3].text = str(d.get("observado", 0)); row[4].text = str(d.get("no_evaluable", 0))
                    for c in row:
                        for pp in c.paragraphs:
                            for rr in pp.runs:
                                rr.font.size = Pt(9)
            doc.add_paragraph("")

        # Controles no conformes
        no_conf = [c for c in cmp.get("controles", []) if c["estado"] == "no_cumple"]
        if no_conf:
            doc.add_heading("Controles no conformes (requieren remediación)", level=2)
            for c in no_conf:
                p = doc.add_paragraph()
                hr = p.add_run(f"{c['instrumento']} {c['referencia']} — {c['titulo']} ")
                hr.bold = True; hr.font.size = Pt(11)
                sr = p.add_run(f"(severidad {c['severidad']})")
                sr.font.size = Pt(10)
                sr.font.color.rgb = RGBColor.from_string(_SEV_COLOR.get(c["severidad"], "1A1A1A"))
                pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pe.add_run("Exigencia: ").bold = True; pe.add_run(c["exigencia"])
                if c.get("hallazgos"):
                    refs = "; ".join(
                        f"{h['rule_id']} en {str(h.get('archivo','')).split('/')[-1]}"
                        + (f":{h['linea']}" if h.get("linea") else "")
                        for h in c["hallazgos"][:5])
                    pv = doc.add_paragraph()
                    pv.add_run(f"Evidencia ({c['n_hallazgos']}): ").bold = True
                    pv.add_run(refs).font.size = Pt(10)

        # Controles no evaluables
        no_eval = [c for c in cmp.get("controles", []) if c["estado"] == "no_evaluable"]
        if no_eval:
            doc.add_heading("Controles no evaluables por análisis estático", level=2)
            doc.add_paragraph(
                "Los siguientes controles son de naturaleza organizacional o procedimental y "
                "deben acreditarse mediante evidencia documental independiente:")
            for c in no_eval:
                p = doc.add_paragraph(style="List Bullet"); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(f"{c['instrumento']} {c['referencia']} — {c['titulo']}. ").bold = True
                if c.get("evidencia_documental"):
                    p.add_run(f"Evidencia esperada: {c['evidencia_documental']}").font.size = Pt(10)

    # ---------- 3. PANORAMA TÉCNICO ----------
    doc.add_heading("3. Panorama técnico del análisis estático", level=1)
    doc.add_paragraph(f"Total de observaciones: {rs.get('total',0)}. Nivel de riesgo agregado: "
                      f"{an.get('nivel_global','—')}.")
    sev = rs.get("por_severidad", {})
    p = doc.add_paragraph("Por severidad: ")
    for k in ("crítico", "alto", "medio", "bajo", "info"):
        if k in sev:
            run = p.add_run(f"{k}: {sev[k]}   "); run.bold = True
            run.font.color.rgb = RGBColor.from_string(_SEV_COLOR[k])
    doc.add_paragraph("Por ámbito: " + ", ".join(f"{k}: {v}" for k, v in rs.get("por_categoria", {}).items()))
    if st.get("lenguajes"):
        doc.add_paragraph("Lenguajes: " + ", ".join(f"{k} ({v})" for k, v in st["lenguajes"].items()))
    if inv.get("agregados"):
        doc.add_paragraph(f"Componentes integrados automáticamente al inventario de software: "
                          f"{inv['agregados']}.")
    if st.get("permisos_android"):
        doc.add_paragraph("Permisos Android declarados: " + ", ".join(
            p.split(".")[-1] for p in st["permisos_android"]))
    if an.get("hotspots"):
        doc.add_paragraph("Puntos calientes (archivos con múltiples hallazgos): " + "; ".join(
            f"{h['archivo'].split('/')[-1]} ({len(h['reglas'])})" for h in an["hotspots"][:6]))
    # Motores de análisis (defensa en profundidad) y flujo de datos.
    mot = (R.get("motores") or {}).get("por_motor") or rs.get("por_motor") or {}
    if mot:
        doc.add_paragraph("Motores de análisis (hallazgos por motor): "
                          + ", ".join(f"{k}: {v}" for k, v in mot.items()))
    corro = [f for f in R.get("findings", []) if f.get("corroborado_por")]
    if corro:
        doc.add_paragraph(f"Hallazgos corroborados por más de un motor: {len(corro)} (mayor confianza).")
    taint = R.get("_taint") or []
    if taint:
        doc.add_paragraph(f"Rutas de flujo de datos no confiable (fuente → sumidero): {len(taint)}.")
        for tp in taint[:6]:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(f"{tp['archivo'].split('/')[-1]}: {tp['fuente']} (L{tp['fuente_linea']}) → "
                      f"{tp['sumidero']} (L{tp['sumidero_linea']}) · {tp.get('cwe','')}").font.size = Pt(10)

    # ---------- 4. ANÁLISIS TÉCNICO DETALLADO ----------
    doc.add_heading("4. Análisis técnico detallado", level=1)
    intro = doc.add_paragraph(
        "Esta sección desarrolla cada observación en el contexto real de uso del código "
        "(inteligencia local de la herramienta). El nivel de riesgo se calcula como impacto × "
        "probabilidad, justificando ambos factores. El mapeo a la normativa y la ley se trata "
        "por separado en la sección 2.")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _campo(label, val, size=11):
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lr = pp.add_run(f"{label} "); lr.bold = True; lr.font.size = Pt(size)
        vr = pp.add_run(val); vr.font.size = Pt(size)
        return pp

    for i, f in enumerate(_cs_sorted(R.get("findings", [])), 1):
        ctx = f.get("contexto") or {}
        nr = f.get("nivel_riesgo", "—")
        nrcol = _NR_COLOR.get(nr, "1A1A1A")
        h = doc.add_paragraph()
        _cs_topborder(h)  # separador de sección sin línea en blanco
        hr = h.add_run(f"4.{i} [{nr}] {f['titulo']} ({f['rule_id']}"
                       + (f" · {f['cwe']}" if f.get("cwe") else "") + ")")
        hr.bold = True; hr.font.size = Pt(12)
        hr.font.color.rgb = RGBColor.from_string(nrcol)

        if f.get("descripcion_contextual"):
            _campo("Descripción (contexto detectado):", f["descripcion_contextual"])
        if f.get("explicacion_no_tecnica"):
            _campo("Qué significa, en simple:", f["explicacion_no_tecnica"])

        loc = f"{f['archivo']}" + (f":{f['linea']}" if f.get("linea") else "")
        ubic = (f"{loc} · rol del archivo: {ctx.get('rol_archivo','—')}"
                + (f" · {ctx.get('constructo')}" if ctx.get("constructo") else "")
                + f" · exposición: {f.get('exposicion','—')}")
        _campo("Ubicación:", ubic)

        # Procedencia (motor) y corroboración.
        motor = f.get("motor", "interno")
        proc = f"motor: {motor}"
        if f.get("corroborado_por"):
            proc += f"; corroborado por: {', '.join(f['corroborado_por'])}"
        if f.get("secreto_prob") is not None:
            proc += (f"; verosimilitud de secreto (ML): {f['secreto_prob']:.0%}"
                     + (" — posible falso positivo" if f.get("posible_falso_positivo") else ""))
        _campo("Procedencia y verificación:", proc + ".")
        # Explotabilidad por flujo de datos (taint).
        if f.get("alcanzable_taint") and f.get("ruta_flujo"):
            rf = f["ruta_flujo"]
            pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            er = pe.add_run("Explotabilidad (flujo de datos): "); er.bold = True
            er.font.color.rgb = RGBColor.from_string(nrcol)
            pe.add_run(f"ruta detectada {rf['fuente']} (línea {rf['fuente_linea']}) → "
                       f"{rf['sumidero']} (línea {rf['sumidero_linea']}) vía «{rf['var']}»; "
                       "el dato no confiable alcanza la operación sensible, por lo que la "
                       "observación se considera potencialmente explotable.")
        frag = ctx.get("fragmento") or []
        if frag:
            le = doc.add_paragraph(); le.add_run("Evidencia en contexto:").bold = True
            for (ln, txt, es) in frag:
                pf = doc.add_paragraph()
                rr = pf.add_run(f"{'▶' if es else '  '} {ln:>4} | {txt}")
                rr.font.name = "Consolas"; rr.font.size = Pt(9)
                if es:
                    rr.bold = True
                    rr.font.color.rgb = RGBColor.from_string(nrcol)
                else:
                    rr.font.color.rgb = RGBColor.from_string("555555")
        else:
            _campo("Evidencia (extracto):", f.get("evidencia", "—"))

        # Cálculo del nivel de riesgo (impacto × probabilidad) con justificación.
        rp = doc.add_paragraph()
        rp.add_run("Cálculo del nivel de riesgo. ").bold = True
        rp.add_run(f"Severidad intrínseca de la regla: {f.get('severidad','—')}.").font.size = Pt(11)
        _campo("· Impacto:", f"{f.get('impacto_nivel','—')} ({f.get('impacto','?')}/5). "
               f"{f.get('impacto_just','')}")
        _campo("· Probabilidad:", f"{f.get('probabilidad_nivel','—')} ({f.get('probabilidad','?')}/5). "
               f"{f.get('probabilidad_just','')}")
        res = doc.add_paragraph()
        res.add_run("· Resultado: ").bold = True
        rr = res.add_run(f"{f.get('impacto','?')} × {f.get('probabilidad','?')} = "
                         f"{f.get('riesgo_valor','?')} → nivel de riesgo {nr}.")
        rr.bold = True; rr.font.color.rgb = RGBColor.from_string(nrcol)

        if f.get("impacto_contexto"):
            _campo("Impacto en el contexto del código:", f["impacto_contexto"])
        _campo("Mitigación puntual:", f.get("mitigacion_puntual", f.get("mitigacion", "")))
        if f.get("controles"):
            tp = doc.add_paragraph()
            tr = tp.add_run(f"Trazabilidad de cumplimiento (ver sección 2): "
                            f"{', '.join(f['controles'])}.")
            tr.italic = True; tr.font.size = Pt(9)
            tr.font.color.rgb = RGBColor.from_string("555555")

    # ---------- 5. MARCO NORMATIVO ----------
    if cmp.get("instrumentos"):
        doc.add_heading("5. Marco normativo de referencia", level=1)
        for inst in cmp["instrumentos"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{inst['id']} — ").bold = True
            p.add_run(f"{inst['titulo']} ({inst['ley']}, {inst['fecha']}). {inst['url']}")

    # ---------- 6. ANEXO METODOLÓGICO ----------
    doc.add_heading("6. Anexo metodológico", level=1)
    p = doc.add_paragraph(
        "El análisis se realiza mediante inspección estática de patrones (SAST ligero) sobre el "
        "código fuente extraído en un entorno aislado, sin ejecutarlo. Cada regla mapea a una "
        "cláusula de la Guía SEGPRES, la norma legal aplicable, el control normativo auditable y, "
        "cuando corresponde, el identificador CWE. El motor de cumplimiento cruza los hallazgos "
        "con un catálogo de controles derivados artículo por artículo de la normativa, emitiendo "
        "por control uno de cuatro estados: cumple, no cumple, observado o no evaluable. Los "
        "controles organizacionales se reportan como no evaluables con indicación de la evidencia "
        "documental que los acreditaría.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2 = doc.add_paragraph()
    p2.add_run("Análisis avanzado (robustecimiento). ").bold = True
    p2.add_run(
        "El alcance se refuerza con tres técnicas complementarias: (1) motores SAST de código "
        "abierto de la industria —Semgrep (multi-lenguaje, con reglas locales offline), "
        "detect-secrets y Bandit— cuyos hallazgos se normalizan e integran; cuando dos motores "
        "coinciden, el hallazgo se marca como corroborado (mayor confianza). (2) Análisis de "
        "flujo de datos (taint) intra-archivo, que rastrea entradas no confiables (fuentes) hasta "
        "operaciones sensibles (sumideros) para identificar rutas potencialmente explotables; "
        "esto aproxima el razonamiento de ejecución sin ejecutar el código. (3) Un clasificador "
        "de verosimilitud de secretos (regresión logística) que ajusta la probabilidad y "
        "distingue secretos reales de marcadores de posición.")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3 = doc.add_paragraph()
    p3.add_run("Sobre las pruebas dinámicas (DAST). ").bold = True
    p3.add_run(
        "Por seguridad, la herramienta NUNCA ejecuta el código auditado: ejecutar artefactos no "
        "confiables expondría al evaluador. Por ello no se realiza DAST en sentido estricto "
        "(ejecución del sistema en marcha); el componente «dinámico» se cubre mediante el análisis "
        "de flujo de datos y el razonamiento de rutas de ataque asistido por IA. Para pruebas "
        "dinámicas completas (con la aplicación desplegada en un entorno controlado) se recomienda "
        "complementar con un proceso de DAST dedicado.")
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if cmp.get("disclaimer"):
        pd = doc.add_paragraph(); dr = pd.add_run(cmp["disclaimer"])
        dr.italic = True; dr.font.size = Pt(9); dr.font.color.rgb = RGBColor.from_string("444444")

    buf = io.BytesIO(); doc.save(buf)
    return _fix_docx_zoom(buf.getvalue())


def codescan_to_pdf(report, author="no indicado", generated_at=""):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    R = report
    rs = R.get("resumen", {}); st = R.get("stats", {}); an = R.get("analisis", {})
    arch = R.get("archive", {}); inv = R.get("inventario", {})
    cmp = R.get("cumplimiento") or {}
    buf = io.BytesIO()
    docp = SimpleDocTemplate(buf, pagesize=letter, topMargin=inch, bottomMargin=inch,
                             leftMargin=inch, rightMargin=inch,
                             title="Informe de auditoría de cumplimiento")
    # Interlineado 1.0: leading ≈ fontSize; sin espacio entre párrafos (spaceAfter=0).
    body = ParagraphStyle("body", fontName="Helvetica", fontSize=10.5, leading=10.5,
                          alignment=TA_JUSTIFY, textColor=colors.HexColor("#1A1A1A"),
                          spaceAfter=0, spaceBefore=0)
    small = ParagraphStyle("small", parent=body, fontSize=9, leading=9,
                           textColor=colors.HexColor("#333333"), spaceAfter=0)
    mono = ParagraphStyle("mono", fontName="Courier", fontSize=8.5, leading=9.5,
                          textColor=colors.HexColor("#333333"), spaceAfter=0)
    h1 = ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=15, leading=15,
                        textColor=colors.HexColor("#" + _ACCENT), spaceBefore=8, spaceAfter=2)
    h2 = ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12, leading=12,
                        textColor=colors.HexColor("#1A1A1A"), spaceBefore=6, spaceAfter=2)
    h3 = ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=10.5, leading=11,
                        textColor=colors.HexColor("#1A1A1A"), spaceBefore=6, spaceAfter=1)
    title = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=20, leading=22,
                           textColor=colors.HexColor("#" + _ACCENT), spaceAfter=2)

    def esc(s):
        return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = []
    # ---------- PORTADA ----------
    story.append(Paragraph("Informe de auditoría de cumplimiento y aseguramiento de calidad", title))
    story.append(Paragraph("<i>Análisis estático de seguridad de código (SAST) y conformidad "
                           "normativa</i>", small))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"<b>Artefacto analizado:</b> {esc(arch.get('nombre','(s/d)'))} · "
                           f"{arch.get('files','?')} archivos · "
                           f"{round(arch.get('bytes',0)/1048576,1)} MB", small))
    story.append(Paragraph(f"<b>Generado por:</b> {esc(author)} · "
                           f"<b>Fecha y hora:</b> {esc(generated_at)}", small))
    story.append(Paragraph("<b>Clasificación:</b> Uso interno — insumo para evaluación de QA y "
                           "auditoría", small))
    story.append(Paragraph("<b>Marco de referencia:</b> " + esc(_MARCO), small))
    story.append(Paragraph("<i>Análisis estático con razonamiento (aprendizaje automático local y "
                           "modelo de lenguaje cuando está disponible): insumo para QA y auditoría, "
                           "no dictamen jurídico ni certificación. Evidencia de secretos "
                           "enmascarada.</i>", small))
    story.append(Spacer(1, 10))

    # ---------- 1. RESUMEN EJECUTIVO ----------
    story.append(Paragraph("1. Resumen ejecutivo", h1))
    if an.get("resumen_ejecutivo"):
        story.append(Paragraph(esc(an["resumen_ejecutivo"]), body))
    if cmp:
        vc = ("#1E7A34" if cmp.get("veredicto") == "CONFORME" else
              "#B00020" if cmp.get("veredicto") == "NO CONFORME" else "#8A6D00")
        story.append(Paragraph(
            f"<b>Veredicto de cumplimiento normativo:</b> "
            f'<font color="{vc}"><b>{esc(cmp.get("veredicto","—"))}</b></font> '
            f"({cmp.get('pct_cumplimiento',0)}% de controles evaluables conformes). "
            f"{esc(cmp.get('veredicto_detalle',''))}", body))
    story.append(Spacer(1, 6))

    # ---------- 2. CONFORMIDAD NORMATIVA ----------
    if cmp:
        co = cmp.get("conteos", {})
        story.append(Paragraph("2. Conformidad con el marco normativo", h1))
        story.append(Paragraph(
            f"Se evaluaron <b>{cmp.get('total_controles',0)} controles</b> auditables: "
            f"{co.get('cumple',0)} conformes, {co.get('no_cumple',0)} no conformes, "
            f"{co.get('observado',0)} observados y {co.get('no_evaluable',0)} no evaluables "
            f"por análisis estático.", body))

        # Tabla matriz por instrumento
        story.append(Paragraph("Matriz de conformidad por instrumento", h2))
        data = [["Instrumento", "Cumple", "No\ncumple", "Observ.", "No\neval."]]
        for inst in cmp.get("instrumentos", []):
            d = cmp.get("por_instrumento", {}).get(inst["id"], {})
            data.append([Paragraph(f"<b>{esc(inst['id'])}</b> — {esc(inst['titulo'])}", small),
                         str(d.get("cumple", 0)), str(d.get("no_cumple", 0)),
                         str(d.get("observado", 0)), str(d.get("no_evaluable", 0))])
        tbl = Table(data, colWidths=[260, 48, 48, 48, 48])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + _ACCENT)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F7")]),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 8))

        # Función Decreto 7
        pf = cmp.get("por_funcion", {})
        if any(sum(v.values()) for v in pf.values()):
            story.append(Paragraph("Cobertura por función de ciberseguridad (Decreto 7/2023)", h2))
            data = [["Función", "Cumple", "No cumple", "Observado", "No eval."]]
            for fn, d in pf.items():
                if sum(d.values()):
                    data.append([fn, str(d.get("cumple", 0)), str(d.get("no_cumple", 0)),
                                 str(d.get("observado", 0)), str(d.get("no_evaluable", 0))])
            tf = Table(data, colWidths=[140, 73, 73, 73, 73])
            tf.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + _ACCENT)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F7")]),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tf)
            story.append(Spacer(1, 8))

        # Controles no conformes
        no_conf = [c for c in cmp.get("controles", []) if c["estado"] == "no_cumple"]
        if no_conf:
            story.append(Paragraph("Controles no conformes (requieren remediación)", h2))
            for c in no_conf:
                col = _SEV_COLOR.get(c["severidad"], "1A1A1A")
                story.append(Paragraph(
                    f'<font color="#{col}"><b>{esc(c["instrumento"])} {esc(c["referencia"])} — '
                    f'{esc(c["titulo"])}</b></font> <font size="8">(severidad {c["severidad"]})</font>',
                    body))
                story.append(Paragraph(f"<b>Exigencia:</b> {esc(c['exigencia'])}", small))
                if c.get("hallazgos"):
                    refs = "; ".join(
                        f"{h['rule_id']} en {str(h.get('archivo','')).split('/')[-1]}"
                        + (f":{h['linea']}" if h.get("linea") else "")
                        for h in c["hallazgos"][:5])
                    story.append(Paragraph(f"<b>Evidencia ({c['n_hallazgos']}):</b> {esc(refs)}", small))
                story.append(Spacer(1, 3))

        # Controles no evaluables
        no_eval = [c for c in cmp.get("controles", []) if c["estado"] == "no_evaluable"]
        if no_eval:
            story.append(Paragraph("Controles no evaluables por análisis estático", h2))
            story.append(Paragraph("De naturaleza organizacional o procedimental; requieren "
                                   "evidencia documental independiente:", small))
            for c in no_eval:
                doc_ev = f" <i>Evidencia esperada:</i> {esc(c['evidencia_documental'])}" if c.get("evidencia_documental") else ""
                story.append(Paragraph(
                    f"• <b>{esc(c['instrumento'])} {esc(c['referencia'])} — {esc(c['titulo'])}.</b>"
                    + doc_ev, small))
            story.append(Spacer(1, 4))

    # ---------- 3. PANORAMA TÉCNICO ----------
    story.append(Paragraph("3. Panorama técnico del análisis estático", h1))
    sev = rs.get("por_severidad", {})
    story.append(Paragraph(f"<b>{rs.get('total',0)} observaciones</b> · nivel de riesgo agregado: "
                 f"<b>{esc(an.get('nivel_global','—'))}</b> · " +
                 ", ".join(f"{k}: {v}" for k, v in sev.items()), body))
    if st.get("lenguajes"):
        story.append(Paragraph("Lenguajes: " + esc(", ".join(
            f"{k} ({v})" for k, v in st["lenguajes"].items())), small))
    if inv.get("agregados"):
        story.append(Paragraph(f"Componentes integrados al inventario de software: "
                     f"<b>{inv['agregados']}</b>.", small))
    if st.get("permisos_android"):
        story.append(Paragraph("Permisos Android: " + esc(", ".join(
            p.split('.')[-1] for p in st['permisos_android'])), small))
    mot = (R.get("motores") or {}).get("por_motor") or rs.get("por_motor") or {}
    if mot:
        story.append(Paragraph("Motores de análisis (hallazgos por motor): "
                     + esc(", ".join(f"{k}: {v}" for k, v in mot.items())), small))
    corro = [f for f in R.get("findings", []) if f.get("corroborado_por")]
    if corro:
        story.append(Paragraph(f"Hallazgos corroborados por más de un motor: "
                     f"<b>{len(corro)}</b> (mayor confianza).", small))
    taint = R.get("_taint") or []
    if taint:
        story.append(Paragraph(f"Rutas de flujo de datos no confiable (fuente -&gt; sumidero): "
                     f"<b>{len(taint)}</b>.", small))
        for tp in taint[:6]:
            story.append(Paragraph(
                f"&nbsp;&nbsp;• {esc(tp['archivo'].split('/')[-1])}: {esc(tp['fuente'])} "
                f"(L{tp['fuente_linea']}) -&gt; {esc(tp['sumidero'])} (L{tp['sumidero_linea']}) "
                f"· {esc(tp.get('cwe',''))}", small))
    story.append(Spacer(1, 6))

    # ---------- 4. ANÁLISIS TÉCNICO DETALLADO ----------
    story.append(Paragraph("4. Análisis técnico detallado", h1))
    story.append(Paragraph(
        "Esta sección desarrolla cada observación en el contexto real de uso del código "
        "(inteligencia local de la herramienta). El nivel de riesgo se calcula como impacto × "
        "probabilidad, justificando ambos factores. El mapeo a la normativa y la ley se trata "
        "por separado en la sección 2.", small))
    from reportlab.platypus import HRFlowable
    from reportlab.lib import colors as _rc
    for i, f in enumerate(_cs_sorted(R.get("findings", [])), 1):
        ctx = f.get("contexto") or {}
        nr = f.get("nivel_riesgo", "—")
        col = _NR_COLOR.get(nr, "1A1A1A")
        story.append(HRFlowable(width="100%", thickness=0.5, spaceBefore=6, spaceAfter=2,
                                color=_rc.HexColor("#B0B7C3")))
        story.append(Paragraph(
            f'<font color="#{col}"><b>4.{i} [{nr}] {esc(f["titulo"])}</b></font> '
            f'<font size="8">({f["rule_id"]}'
            + (f' · {esc(f["cwe"])}' if f.get("cwe") else "") + ')</font>', h3))
        if f.get("descripcion_contextual"):
            story.append(Paragraph(f"<b>Descripción (contexto detectado).</b> "
                                   f"{esc(f['descripcion_contextual'])}", body))
        if f.get("explicacion_no_tecnica"):
            story.append(Paragraph(f"<b>Qué significa, en simple.</b> "
                                   f"{esc(f['explicacion_no_tecnica'])}", body))
        loc = esc(f["archivo"]) + (f":{f['linea']}" if f.get("linea") else "")
        story.append(Paragraph(
            f"<b>Ubicación.</b> {loc} · rol del archivo: {esc(ctx.get('rol_archivo','—'))}"
            + (f" · {esc(ctx.get('constructo'))}" if ctx.get("constructo") else "")
            + f" · exposición: {esc(f.get('exposicion','—'))}", small))
        # Procedencia (motor) y corroboración.
        motor = f.get("motor", "interno")
        proc = f"motor: {esc(motor)}"
        if f.get("corroborado_por"):
            proc += f"; corroborado por: {esc(', '.join(f['corroborado_por']))}"
        if f.get("secreto_prob") is not None:
            proc += (f"; verosimilitud de secreto (ML): {f['secreto_prob']:.0%}"
                     + (" — posible falso positivo" if f.get("posible_falso_positivo") else ""))
        story.append(Paragraph(f"<b>Procedencia y verificación.</b> {proc}.", small))
        # Explotabilidad por flujo de datos (taint).
        if f.get("alcanzable_taint") and f.get("ruta_flujo"):
            rf = f["ruta_flujo"]
            story.append(Paragraph(
                f'<font color="#{col}"><b>Explotabilidad (flujo de datos).</b></font> '
                f"ruta detectada: {esc(rf['fuente'])} (línea {rf['fuente_linea']}) -&gt; "
                f"{esc(rf['sumidero'])} (línea {rf['sumidero_linea']}) vía «{esc(rf['var'])}»; el "
                "dato no confiable alcanza la operación sensible, por lo que se considera "
                "potencialmente explotable.", small))
        frag = ctx.get("fragmento") or []
        if frag:
            story.append(Paragraph("<b>Evidencia en contexto.</b>", small))
            for (ln, txt, es) in frag:
                pref = "&#9654;" if es else "&nbsp;&nbsp;"
                line = f"{pref} {ln:>4} | {esc(txt)}"
                if es:
                    line = f'<font color="#{col}"><b>{line}</b></font>'
                story.append(Paragraph(line, mono))
        else:
            story.append(Paragraph(f"<b>Evidencia (extracto).</b> {esc(f.get('evidencia','—'))}", small))
        # Cálculo del nivel de riesgo (impacto × probabilidad) con justificación.
        story.append(Paragraph(f"<b>Cálculo del nivel de riesgo.</b> Severidad intrínseca de la "
                               f"regla: {esc(f.get('severidad','—'))}.", body))
        story.append(Paragraph(f"· <b>Impacto:</b> {f.get('impacto_nivel','—')} "
                               f"({f.get('impacto','?')}/5). {esc(f.get('impacto_just',''))}", small))
        story.append(Paragraph(f"· <b>Probabilidad:</b> {f.get('probabilidad_nivel','—')} "
                               f"({f.get('probabilidad','?')}/5). {esc(f.get('probabilidad_just',''))}", small))
        story.append(Paragraph(
            f'· <b>Resultado:</b> <font color="#{col}"><b>{f.get("impacto","?")} × '
            f'{f.get("probabilidad","?")} = {f.get("riesgo_valor","?")} (sobre 25); nivel de '
            f'riesgo {nr}</b></font>.', small))
        if f.get("impacto_contexto"):
            story.append(Paragraph(f"<b>Impacto en el contexto del código.</b> "
                                   f"{esc(f['impacto_contexto'])}", body))
        story.append(Paragraph(f"<b>Mitigación puntual.</b> "
                               f"{esc(f.get('mitigacion_puntual', f.get('mitigacion','')))}", body))
        if f.get("controles"):
            story.append(Paragraph(f"<i>Trazabilidad de cumplimiento (ver sección 2): "
                                   f"{esc(', '.join(f['controles']))}.</i>", small))

    # ---------- 5. MARCO NORMATIVO ----------
    if cmp.get("instrumentos"):
        story.append(Paragraph("5. Marco normativo de referencia", h1))
        for inst in cmp["instrumentos"]:
            story.append(Paragraph(
                f"• <b>{esc(inst['id'])}</b> — {esc(inst['titulo'])} "
                f"({esc(inst['ley'])}, {inst['fecha']}). {esc(inst['url'])}", small))
        story.append(Spacer(1, 6))

    # ---------- 6. ANEXO METODOLÓGICO ----------
    story.append(Paragraph("6. Anexo metodológico", h1))
    story.append(Paragraph(
        "El análisis se realiza mediante inspección estática de patrones (SAST ligero) sobre el "
        "código fuente extraído en un entorno aislado, sin ejecutarlo. Cada regla mapea a una "
        "cláusula de la Guía SEGPRES, la norma legal aplicable, el control normativo auditable y, "
        "cuando corresponde, el identificador CWE. El motor de cumplimiento cruza los hallazgos con "
        "un catálogo de controles derivados artículo por artículo de la normativa, emitiendo por "
        "control uno de cuatro estados: cumple, no cumple, observado o no evaluable.", body))
    story.append(Paragraph(
        "<b>Análisis avanzado (robustecimiento).</b> El alcance se refuerza con: (1) motores SAST "
        "de código abierto —Semgrep (multi-lenguaje, reglas locales offline), detect-secrets y "
        "Bandit— cuyos hallazgos se normalizan e integran; cuando dos motores coinciden, el "
        "hallazgo se marca como corroborado. (2) Análisis de flujo de datos (taint) intra-archivo, "
        "que rastrea entradas no confiables hasta operaciones sensibles para identificar rutas "
        "explotables; aproxima el razonamiento de ejecución sin ejecutar el código. (3) Un "
        "clasificador de verosimilitud de secretos (regresión logística) que ajusta la "
        "probabilidad y distingue secretos reales de marcadores de posición.", body))
    story.append(Paragraph(
        "<b>Sobre las pruebas dinámicas (DAST).</b> Por seguridad, la herramienta nunca ejecuta el "
        "código auditado. No se realiza DAST en sentido estricto (ejecución del sistema en marcha); "
        "el componente «dinámico» se cubre mediante el análisis de flujo de datos y el razonamiento "
        "de rutas de ataque asistido por IA. Para pruebas dinámicas completas se recomienda un "
        "proceso DAST dedicado con la aplicación desplegada en un entorno controlado.", body))
    if cmp.get("disclaimer"):
        story.append(Paragraph(f"<i>{esc(cmp['disclaimer'])}</i>", small))

    docp.build(story)
    return buf.getvalue()
